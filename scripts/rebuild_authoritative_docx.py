"""Rebuild the frozen QuantLab 3.0 DOCX publication snapshot in place.

The generated document uses the ``compact_reference_guide`` design preset with
one named typography override for Chinese text (Microsoft YaHei as the East
Asian fallback) and a ``memo_masthead`` opening block.  The root Markdown file
is the authoritative product specification; run this legacy exporter only when
an explicit Word publication snapshot is required, then render and inspect it.
"""

# This source is retained only to reproduce the frozen 3.0 DOCX snapshot.
# Product-semantic changes belong in the root Markdown specification.
# ruff: noqa: E501

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "如何搭建自己的量化交易系统.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def _set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, size: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    style.element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        _set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    _set_style_font(caption, 9, MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def _set_cell_margins(
    cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[idx])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_tr_pr.append(repeat_header)
    header_no_split = OxmlElement("w:cantSplit")
    header_tr_pr.append(header_no_split)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        _set_run_font(run, size=10, color=NAVY, bold=True)
    for values in rows:
        row = table.add_row()
        row_no_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(row_no_split)
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            _set_run_font(run, size=9.5)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _new_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list(doc: Document, items: list[str], *, ordered: bool = False) -> None:
    # Give each logical list its own numbering instance.  Besides resetting
    # ordered lists correctly, this avoids LibreOffice occasionally suppressing
    # a marker when one document-wide bullet sequence crosses page boundaries.
    num_id = _new_numbering(doc, ordered=ordered)
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)
        p_pr.insert(0, num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        _set_run_font(run, size=11)


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        _set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix) :])
        _set_run_font(rest)
    else:
        run = p.add_run(text)
        _set_run_font(run)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_geometry(table, [TABLE_WIDTH_DXA])
    row_no_split = OxmlElement("w:cantSplit")
    table.rows[0]._tr.get_or_add_trPr().append(row_no_split)
    cell = table.cell(0, 0)
    _shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    _set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    _set_run_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    _set_run_font(run, size=9, color=MUTED)


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("QuantLab  ·  产品与验收技术基准")
    _set_run_font(hr, size=8.5, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("第 ")
    _set_run_font(fr, size=9, color=MUTED)
    _field_run(fp, "PAGE")
    fr2 = fp.add_run(" 页")
    _set_run_font(fr2, size=9, color=MUTED)


def _add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("QUANTLAB · 权威产品方案")
    _set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("如何搭建自己的量化交易系统")
    _set_run_font(r, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("个人中低频量化研究、推荐与模拟交易定稿")
    _set_run_font(r, size=14, color=MUTED)

    metadata = [
        ("方案版本", "3.0（定稿）"),
        ("批准日期", "2026-07-16"),
        (
            "适用边界",
            "A股；仅 Tushare；日线与分钟线；研究、回测、推荐、审批与完整模拟交易",
        ),
        ("技术基准", "Qlib + RD-Agent（固定并经兼容验证的版本）"),
        ("唯一排除", "券商连接、真实账户、真实资金与真实委托；QMT 默认关闭"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}：")
        _set_run_font(lr, size=10.5, color=NAVY, bold=True)
        vr = p.add_run(value)
        _set_run_font(vr, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(
        doc,
        "权威声明",
        "本 DOCX 是 3.0 定稿，决定产品目标、个人投资授权、策略体系、数据口径、成本模型、风险参数、推荐与模拟交易规则以及验收标准；Qlib 与 RD-Agent 决定技术实现基准。除非证券市场规则、Tushare 接口或上游软件发生外部变化，否则不得另起一套产品方案。",
    )


def build_document() -> Document:
    doc = Document()
    doc.core_properties.title = "如何搭建自己的量化交易系统"
    doc.core_properties.subject = "个人中低频量化研究、推荐与模拟交易定稿"
    doc.core_properties.author = "QuantLab"
    doc.core_properties.keywords = "Qlib, RD-Agent, Tushare, A股, 回测, 模拟盘"
    doc.core_properties.comments = "Version 3.0 final product, strategy, risk and acceptance baseline"
    doc.core_properties.modified = datetime(2026, 7, 16, tzinfo=UTC)
    _configure_styles(doc)
    _configure_page(doc)
    _add_title_block(doc)

    doc.add_heading("1. 基准、目标与唯一主线", level=1)
    add_paragraph(
        doc,
        "本产品是使用真实市场数据运行的个人投资辅助系统：完成研究、正式回测、买入/卖出/持有推荐、组合风险控制、模拟订单、模拟成交、虚拟持仓和净值跟踪；唯一不执行券商实盘下单。",
    )
    add_table(
        doc,
        ["优先级", "基准", "决定内容"],
        [
            ["1", "本 DOCX", "产品目标、策略、风险、范围和验收"],
            ["2", "Qlib / RD-Agent", "数据表达、研究、模型、组合、回测、执行与实验的技术基准"],
            [
                "3",
                "项目自研",
                "仅补足上游无法满足的 Tushare、治理、A股、持久模拟、配对双腿与运维能力",
            ],
        ],
        [900, 2400, 6060],
    )
    add_callout(
        doc,
        "唯一主线",
        "Tushare 不可变快照 → Qlib 数据集 → RD-Agent 研究 → 独立复算与准入 → Qlib 正式回测 → 策略审批 → 买卖推荐 → 核心/卫星分配 → 统一模拟交易 → 表现复核与策略生命周期",
    )
    doc.add_heading("1.1 个人投资授权", level=2)
    add_table(
        doc,
        ["授权项", "定稿值", "执行含义"],
        [
            ["基准本金", "人民币 500 万元", "回测容量、整手、最低佣金、模拟现金和推荐权重均按此基准"],
            ["投资市场", "A股及境内场内 ETF", "数据源只使用 Tushare"],
            ["频率", "中低频", "核心以日线/月度调仓为主；分钟数据只用于执行模拟和已批准卫星"],
            ["最大组合回撤", "10%", "达到 10% 时停止新增买入推荐并将模拟组合转为退出状态"],
            ["核心目标", "扣费后长期正收益并取得基准超额", "不承诺每笔、每月或每年盈利"],
            ["决策方式", "系统建议、用户决定", "系统输出依据和风险，不替代用户最终判断"],
        ],
        [1900, 2400, 5060],
    )
    add_callout(
        doc,
        "产品目标",
        "系统必须证明策略在锁定样本外、完整成本、不同市场阶段和持续模拟交易中具有正期望，再输出正式推荐。不能满足准入线时，正确结果是保持现金或不推荐，而不是强行给出买卖信号。",
    )
    doc.add_heading("1.2 系统必须输出的结果", level=2)
    add_list(
        doc,
        [
            "每日风险检查；核心策略默认在月度调仓日产生正式买入、卖出、持有和目标权重建议。",
            "每条建议包含证券、动作、当前权重、目标权重、建议执行窗口、原因、风险提示、策略版本和失效条件。",
            "每次推荐同步进入统一模拟交易，记录计划、订单、成交、未成交、现金、持仓、成本、NAV、回撤和相对基准表现。",
            "控制台同时显示回测成绩、样本外成绩、模拟成绩和当前推荐，禁止把回测收益冒充模拟或真实收益。",
        ],
    )
    add_paragraph(
        doc,
        "正式建设顺序固定为指数增强、全市场行业中性多因子、统一推荐与模拟交易，之后才增加卫星策略。卫星必须复用同一准入、回测、审批、分配和模拟链。",
    )
    add_list(
        doc,
        [
            "核心：指数增强。",
            "核心：全市场行业中性多因子。",
            "第一批卫星：Swing 与分钟均值回归。",
            "研究型卫星：配对/统计套利，以及通过准入的 ML 信号。",
        ],
    )

    doc.add_heading("2. 固定边界与禁止事项", level=1)
    add_list(
        doc,
        [
            "数据源只使用 Tushare；权限不足时阻断相关运营验收，不切换其他数据商。",
            "支持日线和分钟线研究、买卖推荐、执行回放和模拟成交，不做 Tick、Level-2、逐笔或毫秒级高频。",
            "除实盘外全部建设；页面、调度器、回测、推荐、回放与模拟任务均不得向 QMT 或任何券商发单。",
            "QMT 仅作为默认关闭的可选插件附录保留，不删除插件，也不进入主线验收。",
            "不建设独立 C++ 回测引擎。",
            "Qlib/RD-Agent 已有等价能力必须迁入、包装或删除重复实现；新旧路径不得同时作为生产路径。",
            "历史策略版本、模拟账户、成交、NAV 与审计记录保持不可变。",
            "任何未经准入的研究结果只能标记为研究候选，不得显示成正式买卖推荐。",
        ],
    )

    doc.add_heading("3. 责任边界与系统架构", level=1)
    add_table(
        doc,
        ["能力域", "技术责任", "项目约束"],
        [
            [
                "数据/特征",
                "Qlib Dataset、表达式、处理器",
                "Tushare 下载、不可变快照、PIT 血缘、Qlib 转换",
            ],
            [
                "自动研究",
                "RD-Agent 候选因子/模型、实验迭代、挑战者编排",
                "独立复算、统计准入、审批和审计",
            ],
            [
                "组合/风险",
                "Qlib PortfolioOptimizer 与 risk model",
                "A股约束、核心/卫星上限、风险状态",
            ],
            [
                "回测/执行",
                "Qlib Strategy、Executor、Exchange、分析与 Recorder",
                "A股 T+1、停牌、涨跌停、费用、流动性",
            ],
            [
                "模拟盘",
                "Qlib 生成目标、订单计划与回测执行证据",
                "PostgreSQL 持久现金、订单、成交、持仓、NAV、事件与幂等",
            ],
            [
                "配对",
                "Qlib 保存 Dataset/Recorder/标准指标",
                "协整/Kalman、卖空资格、借券成本与双腿原子成交",
            ],
            ["运维/UI", "—", "调度、告警、备份恢复及统一控制台"],
        ],
        [1500, 3600, 4260],
    )
    add_paragraph(
        doc,
        "Qlib 与 RD-Agent 必须使用经验证的固定版本，并记录版本与提交标识；任何升级都要重新运行导入、制品兼容、回测结果和量化回归。",
    )

    doc.add_heading("4. 数据、频率与时间语义", level=1)
    add_list(
        doc,
        [
            "Tushare 原始响应按批次形成不可变快照；快照记录接口、参数、请求窗口、时间、校验和与权限证据。",
            "Qlib 数据集只从已冻结且通过质量门的快照生成，数据集记录快照血缘、构建配置、频率、日历和校验和。",
            "原生频率为 day、1min、5min；15min、30min、60min 只由 Qlib 从分钟数据重采样，不增加下载线路。",
            "日线研究和分钟研究使用独立数据集与实验，禁止在同一策略版本中隐式混用。",
            "任何信号最早在下一 Bar 成交；日线信号最早在下一交易日执行，禁止同 Bar 成交和未来信息。",
            "股票池、行业、财务与退市状态必须采用 PIT 口径，包含已退市证券，防止幸存者偏差。",
        ],
    )
    add_table(
        doc,
        ["字段", "含义", "允许值/规则"],
        [
            ["signal_frequency", "生成信号的数据频率", "day / 1min / 5min / Qlib 重采样频率"],
            ["signal_period", "信号观察与持有周期", "策略版本固定，不能静默改变"],
            ["execution_frequency", "执行切片频率", "1min 或 5min；必须晚于信号可用时点"],
            ["execution_contract_hash", "完整执行语义的内容哈希", "回测、回放、模拟批次必须一致"],
        ],
        [2200, 3300, 3860],
    )
    doc.add_heading("4.1 Tushare 字段与公司行动口径", level=2)
    add_table(
        doc,
        ["数据项", "原始口径", "系统标准口径"],
        [
            ["日线 vol", "手", "乘以 100 后存储为股；禁止直接作为 Qlib volume"],
            ["日线 amount", "千元", "乘以 1,000 后存储为人民币元"],
            ["分钟 vol / amount", "股 / 元", "保持股 / 元，不再重复换算"],
            ["OHLC", "未复权价格", "原始价永久保留，执行、涨跌停与账本只使用原始价"],
            ["adj_factor", "复权因子", "因子与收益使用 raw_price × adj_factor；不得使用未来公司行动"],
            ["财务数据", "公告与报告期并存", "最早在实际公告日之后可用；修订按当时可见版本进入新快照"],
            ["行业/指数成分", "纳入、剔除日期", "按交易日 PIT 展开，禁止用当前成分回填历史"],
            ["停牌/退市", "状态与日期", "停牌不成交；退市证券保留在历史股票池并按真实可交易阶段处理"],
        ],
        [1900, 3000, 4460],
    )
    add_list(
        doc,
        [
            "每个数据集必须检查主键重复、交易日缺口、OHLC 关系、负数、零成交量、异常跳变、跨源字段单位和日历一致性。",
            "现金分红、送转、拆并股等公司行动在模拟账本中调整现金或数量；信号计算与模拟记账不得靠同一列复权价同时完成。",
            "停牌或成交量为零时订单保持未成交；一字涨停禁止买入，一字跌停禁止卖出；其他情况按可用分钟 Bar、参与率和成本模型判断。",
            "退市前只使用当时已公开信息；无法成交的退市持仓保持风险标记，按实际现金结算处理，不得用事后可知价格提前退出。",
        ],
    )
    doc.add_heading("4.2 交易成本与冲击模型", level=2)
    add_paragraph(
        doc,
        "以下为 2026-07-16 的模拟基准。税费规则、券商佣金或交易制度变化时，只新增 cost_model_version 并重跑受影响结果，不修改历史记录。",
    )
    add_table(
        doc,
        ["成本项", "定稿基准", "计算规则"],
        [
            ["券商佣金", "成交额 0.03%，单笔最低 5 元", "买卖双边；允许账户配置覆盖，但不得低于真实合同"],
            ["证券交易印花税", "卖出成交额 0.05%", "仅股票卖出；ETF 等按品种规则配置"],
            ["A股交易过户费", "成交额 0.001%", "买卖双边"],
            ["基础滑点", "日线 5 bps；分钟 VWAP 2 bps", "买入加价、卖出减价"],
            ["市场冲击", "25 bps × sqrt(订单额 / 20日平均成交额)", "与基础滑点相加，单边最高 100 bps"],
            ["成交量参与率", "单 Bar 和全日均 ≤1%", "超过部分延迟到下一可执行 Bar，不得强行成交"],
            ["融券情景成本", "年化 8%", "只用于研究型股票多空模拟；不代表真实可借券"],
            ["压力测试", "基础成本 2 倍", "正式准入必须仍满足收益、回撤和容量门槛"],
        ],
        [1900, 2500, 4960],
    )

    doc.add_heading("5. 策略版本、研究和准入", level=1)
    add_paragraph(
        doc,
        "strategy_type 仅保留 multifactor 与 pair；配方包括 index_enhancement、full_market_multifactor、swing_trend、minute_mean_reversion 以及受控 custom。",
    )
    add_paragraph(doc, "每个不可变策略版本必须固化以下内容：")
    add_list(
        doc,
        [
            "配方和配方版本、strategy_type、信号频率、信号周期、执行频率。",
            "风险参数、成本模型、执行窗口、切片、参与率、跨日规则与执行契约哈希。",
            "数据集与不可变快照血缘、Qlib/RD-Agent 版本及提交标识、候选代码哈希。",
            "独立复算结果、Qlib 实验/Recorder 标识、样本内外边界、压力测试与审批证据。",
        ],
    )
    doc.add_heading("5.1 RD-Agent 挑战者流程", level=2)
    add_list(
        doc,
        [
            "RD-Agent 只负责候选因子/模型生成、研究迭代和挑战者实验；不得建立第二套正式回测或自动研究循环。",
            "候选输出不得直接进入组合。项目必须从锁定数据集独立复算因子，检查方向、覆盖、缺失、稳定性和泄露。",
            "挑战者必须使用 Qlib Dataset、Workflow、Recorder 和样本外回测，与不可变基线进行同口径比较。",
            "通过统计、成本、容量、稳健性、血缘和审批门后，候选才可替换或增强基线；否则保留为研究制品。",
        ],
    )
    doc.add_heading("5.2 正式 Qlib 链", level=2)
    add_list(
        doc,
        [
            "风险平价和逆波动率使用 Qlib PortfolioOptimizer；协方差使用 Qlib risk model。",
            "标准绩效指标使用 Qlib 分析组件，年化日数固定为 252，累计收益使用几何口径。",
            "正式回测使用 Qlib Strategy、Executor 和 Exchange；实验通过 Qlib Workflow/Recorder 固化。",
            "A股专用限制通过项目约束包装和 Exchange 扩展实现，不复制 Qlib 的核心求解与分析。",
        ],
    )
    doc.add_heading("5.3 RD-Agent 试验预算与反过拟合", level=2)
    add_list(
        doc,
        [
            "时间切分固定为训练 756 个交易日、验证 252 个交易日、最终测试至少 504 个交易日；所有边界在研究开始前冻结。",
            "每个研究活动最多 100 个候选试验；单个候选最多 30 组参数。所有成功、失败和被丢弃试验均计入试验次数，不得只登记赢家。",
            "最终测试集对同一 experiment_family 只允许解封一次。查看最终测试结果后再调参，必须建立新研究活动并使用更晚的全新测试区间。",
            "多候选比较使用 Benjamini-Hochberg，FDR q ≤0.10；Deflated Sharpe Probability 必须 ≥0.95。",
            "必须通过参数邻域、双倍成本、换手收紧、持仓收紧和零保留缓冲测试；只在单一点参数优秀的候选不得准入。",
            "RD-Agent 不得自行修改准入线、数据边界、成本模型或最终测试集；这些内容只能由新策略版本显式变更。",
        ],
    )
    doc.add_heading("5.4 正式策略准入线", level=2)
    add_paragraph(
        doc,
        "以下是进入正式推荐池的最低条件，不是收益承诺。任一硬门失败，策略只能留在研究池。",
    )
    add_table(
        doc,
        ["指标", "核心策略最低线", "卫星策略最低线"],
        [
            ["样本外长度", "≥504 个交易日", "≥504 个交易日"],
            ["扣费后年化超额收益", "≥3%", "≥4%；无合适基准时扣费后年化收益 ≥6%"],
            ["Sharpe Ratio", "≥0.80", "≥1.00"],
            ["Information Ratio", "≥0.50", "≥0.50；无基准策略可不适用"],
            ["最大回撤", "≤10%", "≤8%"],
            ["滚动稳定性", "3 个 252 日窗口至少 2 个为正，且最近窗口为正", "同左"],
            ["容量", "500 万元基准下全部约束通过", "500 万元基准下全部约束通过"],
            ["压力测试", "双倍成本后年化收益仍为正", "双倍成本后年化收益仍为正"],
        ],
        [2500, 3430, 3430],
    )

    doc.add_heading("6. 核心策略一：指数增强", level=1)
    add_paragraph(
        doc,
        "以 Qlib 六因子基线为不可变起点，RD-Agent 挑战者只能在独立复算和样本外准入后替换或增强基线。",
    )
    add_table(
        doc,
        ["因子组", "基线权重", "处理要求"],
        [
            ["动量", "20%", "PIT、缩尾、z-score、行业/市值中性化"],
            ["反转", "10%", "同上"],
            ["价值", "20%", "同上"],
            ["质量", "20%", "同上"],
            ["成长", "10%", "同上"],
            ["低波动", "20%", "同上"],
        ],
        [2400, 1800, 5160],
    )
    add_table(
        doc,
        ["参数", "默认值", "验收语义"],
        [
            ["基准", "SH000300", "优化与绩效均相对该基准"],
            ["持仓数", "100", "按优化后可交易目标计"],
            ["单票上限", "2%", "任何目标和执行后持仓均不得突破"],
            ["行业偏离", "±3%", "相对基准 PIT 行业权重"],
            ["跟踪误差", "≤3%", "同口径年化 ex-ante/ex-post 证据"],
            ["每日换手", "≤15%", "跨日 VWAP 计划同样受限"],
            ["流动性", "20 日平均成交额 ≥5 亿元", "成交量参与率 ≤1%"],
            ["执行", "5min，3 个交易日 VWAP", "10:00–11:20；13:30–14:50"],
        ],
        [2000, 3300, 4060],
    )

    doc.add_heading("7. 核心策略二：全市场行业中性多因子", level=1)
    add_list(
        doc,
        [
            "股票池为 PIT 全 A 股可交易股票；包含历史退市状态，排除当时不可交易证券。",
            "使用与指数增强相同的 Qlib 六因子基线和 RD-Agent 挑战者流程。",
            "SH000300 只用于绩效报告，不作为优化基准。",
            "行业目标按当期全市场流通市值权重计算；规模、价值、成长和波动暴露保持中性。",
        ],
    )
    add_table(
        doc,
        ["参数", "默认值", "约束"],
        [
            ["调仓", "月度", "信号形成后下一可执行 Bar"],
            ["持仓数", "100", "PIT 可交易股票池"],
            ["单票/单行业", "5% / 15%", "目标及执行后同时约束"],
            ["目标波动率", "15%", "Qlib 优化与风险模型"],
            ["流动性", "20 日平均成交额 ≥5 亿元", "参与率 ≤1%"],
            ["执行", "5min，3 个交易日 VWAP", "使用统一执行窗口"],
            ["回撤", "8%", "暂停新增风险，等待复核"],
        ],
        [2200, 3000, 4160],
    )

    doc.add_heading("8. 卫星策略一：配对/统计套利", level=1)
    add_paragraph(
        doc,
        "配对是卫星系统，不是独立主线。保留 Qlib 无法完整表达的配对专用研究与双腿执行能力，同时复用统一数据、Recorder、指标、审批、分配和模拟接口。",
    )
    add_list(
        doc,
        [
            "研究包含相关性、Engle–Granger 协整、Kalman 动态对冲比率、Z-score、容量、费用、借券与鲁棒性验证。",
            "日线形成信号，下一交易日使用 1min 执行窗口；禁止同日同 Bar 偷看成交。",
            "Tushare 只能提供公开融资融券标的等资格线索，不能证明个人券商账户的真实可借数量和费率；股票空腿统一标记为假设性融券模拟。",
            "双腿订单共享 atomic_group_id，并记录 leg_no、position_side 和 borrow_cost；两腿全部成交或全部拒绝。",
            "配对页面只提供研究、审批和模拟，不提供实盘或券商发单。",
            "没有人工提供的可借数量和费率证据时，股票多空配对不得进入正式投资推荐，只能进入研究报告；ETF 或 long-long 相对价值可按普通卫星准入。",
        ],
    )

    doc.add_heading("9. 其他卫星", level=1)
    doc.add_heading("9.1 Swing", level=2)
    add_paragraph(
        doc,
        "Swing 使用 Qlib 策略主线和现有 MA、ADX、量能、波动、质量过滤配方。默认止损 7%，盈利 12% 减半、20% 清仓；成员回撤 10% 降仓、15% 清仓。",
    )
    doc.add_heading("9.2 分钟均值回归", level=2)
    add_paragraph(
        doc,
        "第一版为多头超跌回归，默认 5min 信号、1min 或 5min 执行，股票遵守 T+1。研究数据、执行数据和模拟批次必须共享频率与契约血缘。",
    )
    doc.add_heading("9.3 ML、行业轮动与事件驱动", level=2)
    add_list(
        doc,
        [
            "ML 只作为 RD-Agent 因子或模型来源，不建设独立 ML 引擎。",
            "行业轮动和事件驱动本轮不建设独立系统；未来若批准，仍必须进入同一 Qlib/RD-Agent 主线。",
        ],
    )

    doc.add_heading("10. 核心/卫星分配与组合风险", level=1)
    add_paragraph(
        doc,
        "分配成员必须记录 role=core|satellite、风险预算和成员上限；数值优化由 Qlib 负责，项目只实施产品级上限、现金、相关性与风险状态。",
    )
    add_table(
        doc,
        ["风险规则", "默认值", "处置"],
        [
            ["核心总权重", "≥70%", "不满足则拒绝审批"],
            ["卫星总权重", "≤30%", "不满足则拒绝审批"],
            ["单卫星", "≤15%", "按成员上限和风险预算共同限制"],
            ["单策略", "≤70%", "核心与卫星均适用"],
            ["目标波动率", "15%", "Qlib PortfolioOptimizer 求解"],
            ["相关性", "≤0.70", "超过门槛阻断新增成员/风险"],
            ["成员回撤", "6% / 8% / 10%", "分别告警 / 暂停新增风险 / 退出并复核"],
            ["组合回撤", "6% / 8% / 10%", "分别降至70% / 降至30% / 停止买入并退出风险持仓"],
        ],
        [2400, 1800, 5160],
    )
    add_paragraph(
        doc,
        "风险平价、逆波动率和目标波动率使用 Qlib 求解；分配审批只能使用认证模拟 NAV，不得用回测收益冒充模拟证据。",
    )
    doc.add_heading("10.1 暂停、恢复与淘汰", level=2)
    add_list(
        doc,
        [
            "任何数据血缘、未来信息、成本、账本、重复记账或指标口径错误，立即暂停受影响策略和推荐，不等待收益表现确认。",
            "成员回撤达到 8%、最近 60 个交易日超额收益 ≤-5%、最近 60 日 IR <0，或连续两个滚动窗口未达准入线时，暂停新增买入推荐。",
            "暂停后必须使用冻结版本复盘；只有数据和代码问题已修复、重新回测通过，并连续 20 个模拟交易日恢复正常，才允许重新启用。",
            "成员回撤达到 10%、连续 120 个交易日未恢复、连续两个季度未达最低线，或策略经济逻辑已失效时，策略版本进入 retired，不得恢复，只能以新版本重新准入。",
            "组合回撤达到 10% 时，模拟账户停止新增买入、生成退出建议并进入人工复核；系统不得为了维持持仓而放宽风险阈值。",
        ],
    )

    doc.add_heading("11. 统一模拟盘", level=1)
    doc.add_heading("11.1 正式推荐", level=2)
    add_paragraph(
        doc,
        "日线策略在 T 日收盘且数据质量门通过后生成 T+1 推荐；系统不得使用 T+1 才公布的数据改写 T 日建议。默认只对已审批且未暂停的策略生成正式推荐。",
    )
    add_table(
        doc,
        ["字段", "必填内容", "用途"],
        [
            ["action", "BUY / SELL / HOLD / EXIT / NO_ACTION", "明确动作；无合格机会时必须允许 NO_ACTION"],
            ["current_weight / target_weight", "当前与目标权重", "计算建议交易量和组合风险"],
            ["quantity_hint", "按500万元和整手规则计算", "仅为人工参考，不连接券商"],
            ["execution_window", "建议日期、时段、最长延迟", "模拟成交与人工执行参考"],
            ["reason", "因子、模型、风险或退出原因", "可解释与复核"],
            ["risk_flags", "停牌、涨跌停、流动性、回撤、数据状态", "任何硬阻断必须显示"],
            ["lineage", "数据集、策略版本、Recorder、成本与执行契约", "保证推荐可复算"],
            ["valid_until", "失效时间或失效条件", "过期建议不得继续显示为有效"],
        ],
        [2300, 3530, 3530],
    )
    doc.add_heading("11.2 模拟订单、成交与账本", level=2)
    add_list(
        doc,
        [
            "只使用 simulation_* 持久账本，不恢复 paper_* 或 pair_paper_* 为活跃生产路径。",
            "主模拟账户初始现金固定为 500 万元；另可建立压力账户，但不得替代主账户验收。",
            "账户以 source_type + source_id + execution_dataset 绑定推荐组合、已审批策略版本或已审批分配版本；旧 recommendation 账户自动映射。",
            "Qlib 负责生成目标、订单计划和回测执行证据；PostgreSQL 负责现金、订单、成交、持仓、NAV、事件、重启恢复和审计。",
            "长仓执行适配器支持 1min/5min、T+1、停牌、涨跌停、税费、流动性和参与率。",
            "配对适配器支持卖空资格、借券成本和双腿原子成交；任一腿不可执行时整组拒绝。",
            "策略、Qlib 回测、执行回放和模拟批次共享 execution_contract_hash；不一致时 fail closed。",
            "批次、订单、成交和 NAV 具有稳定幂等键；重启和重试不得重复记账。",
            "未成交订单必须记录原因并按策略规则取消或延期；禁止在没有可成交 Bar 时按收盘价虚构成交。",
            "每日收盘后对账现金、持仓、成本、公司行动、NAV、回撤和基准收益；任何不平衡都阻断下一批推荐。",
        ],
    )
    add_table(
        doc,
        ["模拟来源", "前置条件", "执行适配器"],
        [
            ["recommendation", "推荐组合有效且血缘一致", "long"],
            ["strategy_version", "策略版本已审批，Qlib 证据完整", "long 或 pair"],
            ["allocation_version", "分配已审批，成员模拟 NAV 已认证", "按成员路由，统一账户记账"],
        ],
        [2300, 4100, 2960],
    )

    doc.add_heading("12. 审批、审计与 readiness", level=1)
    add_list(
        doc,
        [
            "策略、因子、分配和模拟操作均执行认证、权限检查、审批状态机和不可变审计。",
            "readiness 汇总数据血缘、Qlib/RD-Agent 版本、测试/制品、策略状态、模拟证据、风险事件、过期批次和恢复能力。",
            "任何严重风险、过期批次、血缘不一致、频率混用、契约哈希不一致或 Tushare 权限缺口都必须显式阻断相应阶段。",
            "删除重复实现之前必须完成新旧结果回归与生产引用扫描；删除后扫描确认不存在第二条活跃研究、回测、优化或模拟路径。",
        ],
    )

    doc.add_heading("13. 调度、告警、备份恢复与界面", level=1)
    add_list(
        doc,
        [
            "连续研究、参数实验、数据构建、Qlib 回测、模拟回放和分配刷新由统一调度器编排，具备重试、取消、超时和幂等。",
            "告警覆盖数据缺口、血缘、准入失败、回测异常、执行阻断、回撤状态、过期批次、备份和恢复演练。",
            "备份包含 PostgreSQL、不可变快照索引、Qlib 数据集/Recorder 制品和必要配置；恢复后重新校验哈希与 readiness。",
            "控制台统一为数据、RD-Agent、因子准入、Qlib 回测、审批、核心/卫星分配和模拟盘。",
            "配对入口只暴露研究、审批和模拟；QMT/券商入口不得从主导航、调度和模拟任务触发。",
        ],
    )

    doc.add_heading("14. 公共接口与迁移规则", level=1)
    add_table(
        doc,
        ["对象", "新增/固定字段", "迁移原则"],
        [
            ["策略配方", "full_market_multifactor", "旧配方不变，新契约用新版本"],
            ["组合方式", "industry_neutral_qp", "Qlib 优先，项目包装 A股约束"],
            [
                "策略版本",
                "signal_frequency、signal_period、execution_frequency、execution_contract_hash",
                "历史行不可变",
            ],
            ["上游适配器", "Qlib/RD-Agent version + commit", "升级触发兼容和量化回归"],
            ["分配成员", "role、risk_budget、member_cap", "旧记录按明确规则映射并保留审计"],
            ["模拟账户", "source_type、source_id、execution_dataset", "旧 recommendation 自动映射"],
            [
                "订单/成交/持仓",
                "atomic_group_id、leg_no、position_side、borrow_cost",
                "配对原子组不可部分落账",
            ],
        ],
        [1900, 4300, 3160],
    )

    doc.add_heading("15. 测试与验收", level=1)
    doc.add_heading("15.1 自动化覆盖", level=2)
    add_list(
        doc,
        [
            "文档白名单、唯一 DOCX、OOXML 有效性、本地链接和退役路径扫描。",
            "Qlib/RD-Agent 固定版本、导入、运行、Recorder/制品兼容，以及仓库无第二条活跃生产路径。",
            "PIT 股票池、行业、财务和退市股票；因子、频率、重采样、下一 Bar 执行和反泄露。",
            "Qlib 优化器、风险模型、指标口径；累计收益几何口径和年化 252。",
            "TWAP/VWAP 窗口、切片、跨日、费用、参与率、停牌、涨跌停和 T+1。",
            "回撤状态、Swing 退出、配对双腿原子成交、卖空资格和借券成本。",
            "回测、回放与模拟 execution_contract_hash 一致；重启/重试幂等；不产生 QMT/券商请求。",
        ],
    )
    doc.add_heading("15.2 量化与运营完成标准", level=2)
    add_list(
        doc,
        [
            "两个核心策略分别至少 504 个样本外交易日。",
            "至少 3 个滚动 252 日窗口和 5 个最差 20 日事件窗口。",
            "通过双倍成本、收紧换手、持仓和容量压力测试。",
            "形成两个已审批核心版本和一个已审批核心分配；配对卫星不是首期定版验收的阻塞项。",
            "每个核心成员及组合至少完成连续 60 个交易日的前向模拟；系统稳定性验收覆盖至少 120 个交易日回放。",
            "前向模拟期间不得修改历史推荐、成交、成本或 NAV；与同版本回放不一致时必须阻断。",
            "不存在未解决严重风险、过期批次或血缘不一致。",
            "Tushare 权限不足时可完成代码验收，但运营验收必须保持阻断。",
        ],
    )
    doc.add_heading("16. 固定实施顺序", level=1)
    add_table(
        doc,
        ["阶段", "必须完成", "进入下一阶段的条件"],
        [
            ["第一阶段", "Tushare日线/PIT数据、Qlib数据集、指数增强、成本与正式回测", "指数增强通过全部准入线"],
            ["第二阶段", "全市场行业中性多因子、组合风险与正式推荐", "两个核心均通过并可生成可复算推荐"],
            ["第三阶段", "500万元统一模拟账户、订单成交、NAV、回撤、告警、恢复", "连续60日模拟和120日回放通过"],
            ["第四阶段", "Swing、分钟均值回归和ML挑战者", "逐个准入，不影响核心主线"],
            ["第五阶段", "配对/统计套利研究与假设性融券模拟", "有明确借券假设；股票空腿默认不进入正式推荐"],
        ],
        [1500, 4300, 3560],
    )
    add_callout(
        doc,
        "定版完成定义",
        "当第三阶段验收通过时，系统已经是一套可真实运行、能辅助个人投资的中低频量化系统：它使用真实数据生成正式建议，并用完整模拟交易持续验证建议质量。第四、第五阶段属于能力扩展，不改变主系统是否成立。",
    )

    doc.add_heading("附录 A：QMT 可选插件边界", level=1)
    add_paragraph(
        doc,
        "QMT/券商网关保留为默认关闭的可选插件，不属于当前研究、回测、审批或模拟主线。启用插件需要另行产品批准、独立安全评审和验收；在此之前：",
    )
    add_list(
        doc,
        [
            "BROKER_FEATURE_ENABLED 默认且必须为 false。",
            "主页面、计划任务、模拟批次和回放任务不得调用网关。",
            "任何实盘环境、真实账户或真实委托均不受本方案授权。",
        ],
    )

    doc.add_heading("附录 B：外部规则与技术来源", level=1)
    add_list(
        doc,
        [
            "Tushare A股日线行情：https://tushare.pro/document/1?doc_id=27",
            "Tushare 历史分钟行情：https://tushare.pro/document/2?doc_id=370",
            "财政部、税务总局证券交易印花税公告：https://www.mof.gov.cn/jrttts/202308/t20230828_3904235.htm",
            "中国结算上海/深圳市场收费表：https://www.chinaclear.cn/",
            "Qlib 官方仓库与文档：https://github.com/microsoft/qlib",
            "RD-Agent 官方仓库：https://github.com/microsoft/rd-agent",
        ],
    )
    add_paragraph(
        doc,
        "外部链接只用于解释当前参数来源。系统运行时必须把实际采用的接口版本、权限证据、费率生效日和上游提交标识固化到不可变制品中。",
    )

    doc.add_heading("附录 C：修订记录", level=1)
    add_table(
        doc,
        ["版本", "日期", "状态", "修订摘要"],
        [
            [
                "3.0",
                "2026-07-16",
                "定稿",
                "明确500万元个人投资授权和除实盘外的完整范围；固定数据单位、复权与公司行动、费用与冲击、反过拟合、策略准入、推荐输出、模拟账本、暂停淘汰、60日模拟与分阶段验收。",
            ],
            [
                "2.0",
                "2026-07-16",
                "已取代",
                "原位重整为 Qlib + RD-Agent 技术基准单主线；固定 Tushare、非实盘/非高频边界；定义双核心、配对卫星、统一模拟盘、核心/卫星分配与验收。",
            ],
            ["1.x", "历史", "已取代", "原稿由 Git 历史保存，不再作为当前方案。"],
        ],
        [1000, 1500, 1400, 5460],
    )
    return doc


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
